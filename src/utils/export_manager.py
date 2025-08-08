import threading
from tkinter import filedialog, messagebox, Toplevel, ttk
from datetime import datetime
import logging
import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk # For Toplevel and ttk.Progressbar
from src.database.db_manager import DatabaseManager 

def _show_progress_bar(root):
    progress_window = tk.Toplevel(root)
    progress_window.title("Đang xử lý...")
    progress_window.geometry("300x50")
    progress_window.transient(root)
    progress_window.grab_set()
    progress_label = ttk.Label(progress_window, text="Đang xuất file, vui lòng chờ...", padding=5)
    progress_label.pack(side=tk.TOP, fill=tk.X, expand=True)
    progress_bar = ttk.Progressbar(progress_window, mode='indeterminate')
    progress_bar.pack(side=tk.TOP, fill=tk.X, expand=True, padx=10)
    progress_bar.start(10)
    root.update_idletasks()
    logging.getLogger(__name__).info("Progress bar shown.")
    return progress_window

def _hide_progress_bar(progress_window):
    if progress_window:
        progress_window.destroy()
        logging.getLogger(__name__).info("Progress bar hidden.")

def export_excel_report(root, db_path, from_date, to_date, task_filter, status_filter):
    logging.getLogger(__name__).info(f"Starting Excel export for {from_date} to {to_date}.")
    # TẠO KẾT NỐI MỚI TRONG LUỒNG NÀY
    db_manager = DatabaseManager(db_name=db_path)
    
    report_data = db_manager.get_records_by_filters(
        from_date, 
        to_date, 
        task=task_filter,
        status=status_filter
    )
    
    if not report_data:
        db_manager.close() # ĐÓNG KẾT NỐI CỦA LUỒNG NÀY
        root.after(0, lambda: messagebox.showinfo("Cảnh báo", "Không có dữ liệu để xuất Excel."))
        logging.getLogger(__name__).warning("No data found for Excel export.")
        return
    
    filename = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".xlsx",
        filetypes=[("Excel files", "*.xlsx")],
        initialfile=f"BaoCao_CongViec_{datetime.now().strftime('%Y%m%d')}.xlsx"
    )
    
    if not filename:
        db_manager.close() # ĐÓNG KẾT NỐI CỦA LUỒNG NÀY
        logging.getLogger(__name__).info("Excel export cancelled by user.")
        return
    
    progress_window = _show_progress_bar(root)

    try:
        # ... (phần code tạo và ghi dữ liệu vào file Excel không thay đổi)
        wb = Workbook()
        ws = wb.active
        ws.title = "Báo cáo công việc"
        
        title_font = Font(bold=True, size=16, color="000000")
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                        top=Side(style='thin'), bottom=Side(style='thin'))
        
        ws.merge_cells('A1:E1')
        title_cell = ws['A1']
        title_cell.value = "BÁO CÁO CÔNG VIỆC"
        title_cell.font = title_font
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        
        current_row = 3
        ws[f'A{current_row}'] = f"Thời gian: {from_date} đến {to_date}"
        current_row += 2 
        
        headers = ['Ngày', 'Công việc', 'Phòng/Khoa', 'Chi tiết', 'Trạng thái']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal="center")
        
        for row_data in report_data:
            current_row += 1
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=current_row, column=col_idx, value=value)
                cell.border = border
                if col_idx == 4: # Details column
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
        
        column_widths = [15, 30, 20, 50, 15]
        for col, width in enumerate(column_widths, start=1):
            col_letter = get_column_letter(col)
            ws.column_dimensions[col_letter].width = width
        
        wb.save(filename)
        root.after(0, lambda: messagebox.showinfo("Thành công", f"Đã xuất báo cáo Excel thành công!"))
        logging.getLogger(__name__).info(f"Excel report successfully exported to: {filename}")
    except Exception as e:
        root.after(0, lambda: messagebox.showerror("Lỗi", f"Không thể xuất báo cáo Excel: {e}"))
        logging.getLogger(__name__).error(f"Excel export failed: {e}", exc_info=True)
    finally:
        _hide_progress_bar(progress_window)
        db_manager.close() # ĐẢM BẢO KẾT NỐI LUÔN ĐƯỢC ĐÓNG
def export_word_report(root, db_path, from_date, to_date, task_filter, status_filter):
    logging.getLogger(__name__).info(f"Starting Word export for {from_date} to {to_date}.")

    # TẠO KẾT NỐI MỚI TRONG LUỒNG NÀY
    db_manager = DatabaseManager(db_name=db_path)

    report_data = db_manager.get_records_by_filters(
        from_date, 
        to_date, 
        task=task_filter,
        status=status_filter
    )
    
    if not report_data:
        db_manager.close() # ĐÓNG KẾT NỐI CỦA LUỒNG NÀY
        root.after(0, lambda: messagebox.showinfo("Cảnh báo", "Không có dữ liệu để xuất Word."))
        logging.getLogger(__name__).warning("No data found for Word export.")
        return
    
    filename = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")],
        initialfile=f"BaoCao_CongViec_{datetime.now().strftime('%Y%m%d')}.docx"
    )
    
    if not filename:
        db_manager.close() # ĐÓNG KẾT NỐI CỦA LUỒNG NÀY
        logging.getLogger(__name__).info("Word export cancelled by user.")
        return
    
    progress_window = _show_progress_bar(root)

    try:
        doc = Document()
        
        title = doc.add_heading('BÁO CÁO CÔNG VIỆC', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Thời gian báo cáo: {from_date} đến {to_date}")
        doc.add_paragraph("")
        doc.add_heading('CHI TIẾT CÔNG VIỆC', level=1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        headers = ['Ngày', 'Công việc', 'Phòng/Khoa', 'Chi tiết', 'Trạng thái']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for row_data in report_data:
            row = table.add_row()
            row.cells[0].text = str(row_data[0])
            row.cells[1].text = str(row_data[1])
            row.cells[2].text = str(row_data[2]) if row_data[2] else ""
            row.cells[3].text = str(row_data[3]) if row_data[3] else ""
            row.cells[4].text = str(row_data[4])
        
        doc.save(filename)
        root.after(0, lambda: messagebox.showinfo("Thành công", f"Đã xuất báo cáo Word thành công!"))
        logging.getLogger(__name__).info(f"Word report successfully exported to: {filename}")
    except Exception as e:
        root.after(0, lambda: messagebox.showerror("Lỗi", f"Không thể xuất báo cáo Word: {e}"))
        logging.getLogger(__name__).error(f"Word export failed: {e}", exc_info=True)
    finally:
        _hide_progress_bar(progress_window)
        db_manager.close() # ĐẢM BẢO KẾT NỐI LUÔN ĐƯỢC ĐÓNG